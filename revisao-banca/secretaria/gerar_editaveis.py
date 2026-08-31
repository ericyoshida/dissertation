#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera os quatro arquivos editaveis (.docx) que a Secretaria pede, a partir do
estado atual da dissertacao: folha de rosto, folha da banca, resumo e abstract.

    python3 revisao-banca/secretaria/gerar_editaveis.py

Le tese.tex e os arquivos de PreTextuais, de modo que os .docx nunca fiquem
desatualizados em relacao ao PDF.
"""
import re, pathlib
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

AQUI = pathlib.Path(__file__).resolve().parent
RAIZ = AQUI.parents[1]
tese = (RAIZ / "tese.tex").read_text()


def macro(nome, texto=tese):
    """Argumento de \\nome{...}, respeitando chaves aninhadas."""
    i = texto.find("\\" + nome + "{")
    if i < 0:
        return ""
    i += len(nome) + 2
    d, out = 1, []
    while i < len(texto) and d:
        c = texto[i]
        if c == "{":
            d += 1
        elif c == "}":
            d -= 1
            if not d:
                break
        out.append(c)
        i += 1
    return "".join(out).strip()


def limpa(s):
    s = re.sub(r"(?m)%.*$", "", s)
    s = re.sub(r"\\emph\{([^}]*)\}", lambda m: m.group(1), s)
    s = s.replace("\\'e", "é").replace("--", "–")
    s = re.sub(r"\\[a-zA-Z]+\*?", "", s)
    s = s.replace("{", "").replace("}", "").replace("~", " ")
    return re.sub(r"\s+", " ", s).strip()


def paragrafos(rel):
    t = (RAIZ / rel).read_text()
    t = re.sub(r"(?m)^%.*$", "", t)
    return [limpa(x) for x in re.split(r"\n\s*\n", t) if x.strip()]


TITULO = limpa(macro("title"))
a = re.search(r"\\author\{([^}]*)\}\{([^}]*)\}", tese)
AUTOR = f"{a.group(1)} {a.group(2)}"
CURSO = limpa(macro("course"))
AREA = limpa(macro("area"))
ADV = re.search(r"\\advisor\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}", tese)
COA = re.search(r"\\coadvisor\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}", tese)
ANO = re.search(r"\\date\{\d+\}\{[^}]*\}\{(\d+)\}", tese).group(1)

# A folha de rosto so traz o pro-reitor se \boss for chamado. O template novo do
# ITA nao traz, entao a chamada esta suprimida no tese.tex e este bloco some.
TEM_BOSS = bool(re.search(r"(?m)^\s*\\boss\{", tese))

_exam = [
    (m.group(1), m.group(2), m.group(3), m.group(4))
    for m in re.finditer(r"\\examiner\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}", tese)
]
# A classe insere orientador e coorientador logo apos o presidente, e so os
# demais examinadores depois; a folha da banca impressa segue essa ordem.
banca = _exam[:1] + [
    (ADV.group(1), ADV.group(2), "Advisor", ADV.group(3)),
    (COA.group(1), COA.group(2), "Co-advisor", COA.group(3)),
] + _exam[1:]


def novo():
    d = Document()
    s = d.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(12)
    return d


def par(d, txt, size=12, bold=False, align=AL.CENTER, space=0):
    p = d.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(txt)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold


# ---------------- 1. Folha de rosto ----------------
d = novo()
par(d, f"Dissertation presented to the Instituto Tecnológico de Aeronáutica, in partial "
       f"fulfillment of the requirements for the degree of Master of Science in the Graduate "
       f"Program of {CURSO}, Field of {AREA}.", 14, align=AL.JUSTIFY, space=36)
par(d, AUTOR, 14, True, space=36)
par(d, TITULO.upper(), 16, True, space=36)
par(d, "Dissertation approved in its final version by signatories below:", 12, space=36)
par(d, f"{ADV.group(1).replace('~', ' ')} {ADV.group(2)}", 14)
par(d, "Advisor", 14, space=30)
par(d, f"{COA.group(1).replace('~', ' ')} {COA.group(2)}", 14)
par(d, "Co-advisor", 14, space=30)
if TEM_BOSS:
    b = re.search(r"(?m)^\s*\\boss\{([^}]*)\}\{([^}]*)\}", tese)
    par(d, f"{b.group(1).replace('~', ' ')} {b.group(2)}", 14)
    par(d, "Pro-Rector of Graduate Courses", 14, space=30)
par(d, "Campo Montenegro", 12)
par(d, "São José dos Campos, SP – Brazil", 12)
par(d, ANO, 12)
d.save(AQUI / "EDITAVEL_1_Folha_de_rosto.docx")

# ---------------- 2. Folha da banca ----------------
d = novo()
par(d, TITULO.upper(), 16, True, space=36)
par(d, AUTOR, 14, True, space=36)
par(d, "Thesis Committee Composition:", 12, align=AL.LEFT, space=18)
t = d.add_table(rows=0, cols=4)
for tit, nome, papel, inst in banca:
    c = t.add_row().cells
    for cell, val in zip(c, (tit.replace("~", " "), nome, papel + "  -", inst)):
        cell.text = ""
        r = cell.paragraphs[0].add_run(val)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
d.add_paragraph()
par(d, "ITA", 14, True)
d.save(AQUI / "EDITAVEL_2_Folha_da_banca.docx")

# ---------------- 3 e 4. Resumo e Abstract ----------------
for rel, titulo, saida in (
    ("PreTextuais/resumo.tex", "RESUMO", "EDITAVEL_3_Resumo.docx"),
    ("PreTextuais/abstract.tex", "ABSTRACT", "EDITAVEL_4_Abstract.docx"),
):
    d = novo()
    par(d, titulo, 14, True, space=18)
    for p in paragrafos(rel):
        par(d, p, 12, align=AL.JUSTIFY, space=12)
    d.save(AQUI / saida)

print("gerados a partir do tese.tex:")
print("   folha de rosto  (pro-reitor:",
      "presente" if TEM_BOSS else "ausente, conforme o template novo", ")")
print(f"   folha da banca  ({len(banca)} membros)")
print("   resumo e abstract")
