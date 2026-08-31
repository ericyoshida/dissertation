#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera PREENCHIDO_Folha_de_Registro.docx como transcricao literal da folha de
registro impressa na dissertacao, lendo os mesmos macros de tese.tex.

Rode de novo sempre que tese.tex mudar (por exemplo quando a Biblioteca devolver
o numero de registro) para que os dois documentos continuem identicos:

    python3 revisao-banca/secretaria/gerar_folha_registro.py
"""
import re, pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT

RAIZ = pathlib.Path(__file__).resolve().parents[2]
RED  = RGBColor(0xC0, 0, 0)

def macro(nome, texto):
    """Extrai o argumento de \\nome{...} respeitando chaves aninhadas."""
    i = texto.find("\\" + nome + "{")
    if i < 0: return ""
    i += len(nome) + 2; d = 1; out = []
    while i < len(texto) and d:
        c = texto[i]
        if c == '{': d += 1
        elif c == '}':
            d -= 1
            if not d: break
        out.append(c); i += 1
    return "".join(out).strip()

def limpa(s):
    s = re.sub(r'(?m)%.*$', '', s)
    s = re.sub(r"\\emph\{([^}]*)\}", r"\1", s)
    s = s.replace("\\'e", "é").replace('--', '–')
    s = re.sub(r'\\[a-zA-Z]+\*?', '', s)
    s = s.replace('{', '').replace('}', '').replace('~', ' ')
    return re.sub(r'\s+', ' ', s).strip()

tese = (RAIZ / "tese.tex").read_text()
campos = {
 "data":       limpa(macro("FRDitadata", tese)),
 "docnro":     limpa(macro("FRDitadocnro", tese)),
 "orgao":      limpa(macro("FRDitaorgaointerno", tese)),
 "kw_autor":   limpa(macro("FRDitapalavrasautor", tese)),
 "kw_index":   limpa(macro("FRDitapalavrasresult", tese)),
 "apres":      limpa(macro("FRDitapalavraapresentacao", tese)),
 "titulo":     limpa(macro("title", tese)),
}
autor = re.search(r'\\author\{([^}]*)\}\{([^}]*)\}', tese)
campos["autor"] = f"{autor.group(1)} {autor.group(2)}"
campos["resumo"] = limpa((RAIZ / "PreTextuais/abstract_frd.tex").read_text())

import subprocess
pdf = RAIZ / "tese.pdf"
campos["paginas"] = ""
if pdf.exists():
    txt = subprocess.run(['pdftotext', '-layout', str(pdf), '-'],
                         capture_output=True, text=True).stdout
    frd = next(p for p in txt.split(chr(12))
               if 'FOLHA DE REGISTRO DO DOCUMENTO' in p)
    linhas = [l for l in frd.splitlines() if l.strip()]
    i = next(k for k, l in enumerate(linhas) if 'DE P' in l and 'GINAS' in l)
    nums = re.findall(r'\b(\d+)\b', linhas[i + 1])
    campos["paginas"] = nums[-1] if nums else ""

d = Document()
sec = d.sections[0]
for m in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec, m, Cm(2))
st = d.styles['Normal']; st.font.name = 'Times New Roman'; st.font.size = Pt(10)

def cellset(cell, label, value, red=False):
    cell.text = ""
    p0 = cell.paragraphs[0]; p0.paragraph_format.space_after = Pt(1)
    r = p0.add_run(label); r.font.size = Pt(7.5); r.font.name = 'Times New Roman'
    if value is not None:
        p1 = cell.add_paragraph(); p1.paragraph_format.space_after = Pt(1)
        rv = p1.add_run(value); rv.font.size = Pt(10); rv.font.name = 'Times New Roman'
        if red: rv.font.color.rgb = RED

t0 = d.add_paragraph(); t0.alignment = AL.CENTER
tr = t0.add_run("FOLHA DE REGISTRO DO DOCUMENTO")
tr.font.size = Pt(12); tr.font.name = 'Times New Roman'

t = d.add_table(rows=0, cols=4); t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER

r = t.add_row().cells
cellset(r[0], "1. CLASSIFICAÇÃO/TIPO", "DM")
cellset(r[1], "2. DATA", campos["data"] or " ")
cellset(r[2], "3. DOCUMENTO Nº", campos["docnro"], red=("XXX" in campos["docnro"]))
cellset(r[3], "4. Nº DE PÁGINAS", campos["paginas"])

def full(label, value, red=False):
    row = t.add_row().cells; row[0].merge(row[3])
    cellset(row[0], label, value, red=red)

full("5. TÍTULO E SUBTÍTULO:", campos["titulo"])
full("6. AUTOR(ES):", campos["autor"])
full("7. INSTITUIÇÃO(ÕES)/ÓRGÃO(S) INTERNO(S)/DIVISÃO(ÕES):", campos["orgao"])
full("8. PALAVRAS-CHAVE SUGERIDAS PELO AUTOR:", campos["kw_autor"])
full("9. PALAVRAS-CHAVE RESULTANTES DE INDEXAÇÃO:", campos["kw_index"])
full("10. APRESENTAÇÃO:                                                          "
     "(X) Nacional     ( ) Internacional", campos["apres"])
full("11. RESUMO:", campos["resumo"])
full("12. GRAU DE SIGILO:",
     "(X) OSTENSIVO                    ( ) RESERVADO                    ( ) SECRETO")

saida = pathlib.Path(__file__).with_name("PREENCHIDO_Folha_de_Registro.docx")
d.save(saida)
print("gerado:", saida.name)
for k in ("data", "docnro", "paginas"):
    print(f"  campo {k:8s} = {campos[k]!r}")
