#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Preenche os formularios que levam dados pessoais.

Le revisao-banca/secretaria/.dados_pessoais.json (fora do Git) e escreve em
revisao-banca/secretaria/COM_DADOS_PESSOAIS/ (tambem fora do Git). Nenhum dado
pessoal fica neste arquivo nem no repositorio. Campos ainda vazios no JSON
saem em vermelho como PREENCHER.

    python3 revisao-banca/secretaria/gerar_formularios_pessoais.py
"""
import io, json, pathlib, textwrap
from reportlab.pdfgen import canvas
from reportlab.lib.colors import black, red
from pypdf import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

AQUI = pathlib.Path(__file__).resolve().parent
RAIZ = AQUI.parents[1]
ORIG = AQUI / "originais_em_branco"
SAI  = AQUI / "COM_DADOS_PESSOAIS"; SAI.mkdir(exist_ok=True)
D = json.loads((AQUI / ".dados_pessoais.json").read_text(encoding="utf-8"))
H = 842.0

def cpf_banca(nome):
    """CPF de um membro da banca, ou marcador vermelho se ainda nao informado."""
    x = (D.get("cpfs_banca", {}).get(nome) or "").strip()
    return (x, False) if x else ("PREENCHER: CPF", True)

def v(chave, rotulo=None):
    """Valor do JSON, ou marcador vermelho se ainda vazio."""
    x = (D.get(chave) or "").strip()
    return (x, False) if x else (f"PREENCHER: {rotulo or chave}", True)

def limpa_tex(s):
    import re as _re
    s = _re.sub(r"(?m)%.*$", "", s)
    s = _re.sub(r"\\emph\{([^}]*)\}", lambda m: m.group(1), s)
    s = s.replace("\\'e", "e").replace("--", "-")
    s = _re.sub(r"\\[a-zA-Z]+\*?", "", s)
    s = s.replace("{", "").replace("}", "").replace("~", " ")
    return _re.sub(r"\s+", " ", s).strip()


def overlay(src, dst, itens):
    buf = io.BytesIO(); c = canvas.Canvas(buf, pagesize=(595, 842))
    for it in itens:
        x, y, txt = it[0], it[1], it[2]
        size = it[3] if len(it) > 3 else 9
        col  = red if (len(it) > 4 and it[4]) else black
        c.setFont("Helvetica", size); c.setFillColor(col); c.drawString(x, H - y, txt)
    c.save(); buf.seek(0)
    pg = PdfReader(src).pages[0]; pg.merge_page(PdfReader(buf).pages[0])
    w = PdfWriter(); w.add_page(pg)
    with open(dst, "wb") as f: w.write(f)
    print("  ", dst.name)

TIT = "Object Detection and Tracking Using an Unmanned Aerial Vehicle"
kw  = ("Veiculos aereos nao tripulados; Rastreamento de objetos; "
       "Filtro de Kalman estendido; Redes neurais convolucionais")
print("gerando em", SAI.name + "/")

# ---- CAPES 1 ----
cpf, cpf_r = v("cpf"); mes, mes_r = v("matricula_mes", "mes"); ano, ano_r = v("matricula_ano", "ano")
overlay(ORIG / "Capes1_Identificacao.pdf", SAI / "Capes1_Identificacao.pdf", [
 (115, 120, "Engenharia Eletronica e Computacao (PG/EEC)"),
 (115, 139, D["nome"]), (117, 172, "X"), (370, 172, "Brasil"),
 (115, 191, cpf, 9, cpf_r), (381, 191, mes, 8, mes_r), (424, 191, ano, 8, ano_r),
 (115, 213, "07/2026", 8), (369, 212, "X"),
 (115, 237, TIT, 9),
 (163, 289, "Sistemas Autonomos e Ciencia de Dados", 8),
 # Projeto de Pesquisa: deixado em branco -- o programa vincula o trabalho a um
 # projeto proprio no Sucupira; nao ha projeto vinculado ao aluno.
 (163, 327, "Informatica"), (163, 378, "Biblioteca Central do ITA"),
 (160, 397, "1"), (262, 397, "155"), (380, 397, "Ingles"), (163, 416, kw, 6.5),
 (115, 468, "Marcos Ricardo Omena de Albuquerque Maximo"),
 (117, 486, "X"), (370, 488, "Brasil"), (115, 507, cpf_banca("Marcos Ricardo Omena de Albuquerque Maximo")[0], 9,
  cpf_banca("Marcos Ricardo Omena de Albuquerque Maximo")[1]),
 (115, 527, "Stiven Schwanz Dias"),
 (117, 546, "X"), (370, 547, "Brasil"), (115, 567, cpf_banca("Stiven Schwanz Dias")[0], 9, cpf_banca("Stiven Schwanz Dias")[1]),
 (163, 683, v("financiador")[0], 9, v("financiador")[1]),
 (117, 703, "X" if D.get("financiador_natureza") == "B" else ""),   # Natureza: B - Bolsa
 (518, 705, (D.get("financiador_meses") or "??").strip(), 8,
  not (D.get("financiador_meses") or "").strip()),   # caixa estreita: marcador curto
])

# ---- CAPES 2 ----
banca = [("Paulo Marcelo Tasinaffo",127,147,148,170),
         ("Marcos Ricardo Omena de Albuquerque Maximo",192,213,214,236),
         ("Stiven Schwanz Dias",258,278,279,301),
         ("Marcelo Gomes da Silva Bruno",323,344,345,367),
         ("Alberto Ferreira de Souza",389,409,411,432)]
it2 = []
for nome, yn, ycb, yp, yd in banca:
    doc, doc_r = cpf_banca(nome)
    it2 += [(115, yn, nome), (117, ycb, "X"), (370, yp, "Brasil"),
            (115, yd, doc, 9, doc_r)]
bai, bai_r = v("bairro"); tel, tel_r = v("telefone")
it2 += [
 # dominios oficiais do Manual do Coleta de Dados da CAPES (aba Atividade Futura):
 #  Vinculo: CLT | Servidor Publico | Aposentado | Colaborador | Bolsa de fixacao
 #  Expectativa: Ensino e Pesquisa | Pesquisa | Empresa | Profissional autonomo | Outras
 (175, 556, "CLT"),
 (175, 578, "Empresa"),
 (533, 578, "X"),                       # Mesma Area da Titulacao: Sim

 (115, 635, D["logradouro"]), (115, 657, bai, 9, bai_r), (370, 657, D["cidade"]),
 (115, 679, D["uf"]), (370, 679, "Brasil"), (370, 701, D["cep"]),
 (115, 723, tel, 9, tel_r), (115, 788, D["email"]),
]
overlay(ORIG / "Capes2_Banca.pdf", SAI / "Capes2_Banca.pdf", it2)

# ---- CAPES 3 (codigos da Tabela de Areas do Conhecimento do CNPq) ----
resumo = limpa_tex((RAIZ / "PreTextuais/abstract_frd.tex").read_text())
it3 = [
 (105, 122, "1.03.03.00-6"), (245, 122, "Metodologia e Tecnicas da Computacao"),
 (105, 143, "3.04.02.05-0"), (245, 143, "Sistemas Eletronicos de Medida e de Controle"),
 (105, 165, "3.12.00.00-1"), (245, 165, "Engenharia Aeroespacial"),
]
# pautas do quadro RESUMO vao de 41.8pt a 552.2pt; recuo de ~26pt em cada lado
y = 225.2
for linha in textwrap.wrap(resumo, 118):
    it3.append((68, y, linha, 7.5)); y += 13.68
overlay(ORIG / "Capes3_Area.pdf", SAI / "Capes3_Area.pdf", it3)

# ---- Termo de autorizacao ----
RED = RGBColor(0xC0, 0, 0)
doc = Document(); st = doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
def par(txt, size=12, bold=False, align=AL.CENTER, space=0):
    p = doc.add_paragraph(); p.alignment = align; p.paragraph_format.space_after = Pt(space)
    r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
par("TERMO DE AUTORIZAÇÃO", 14, True, space=2)
par("PARA DISPONIBILIZAÇÃO DE PUBLICAÇÕES", 14, True, space=18)
par("TERMO DE AUTORIZAÇÃO – TESES/DISSERTAÇÕES", 12, True, space=18)
p = doc.add_paragraph(); p.alignment = AL.JUSTIFY
nac,_ = v("nacionalidade"); ec, ec_r = v("estado_civil"); pf, pf_r = v("profissao")
rg, rg_r = v("rg"); org, org_r = v("rg_orgao", "orgao emissor do RG")
end = f"{D['logradouro']}, {bai}, {D['cidade']}, CEP {D['cep']}, {D['uf']}"
for txt, isred in [
  (f"Eu, {D['nome']}, {nac}, ", False), (ec, ec_r), (", ", False), (pf, pf_r),
  (f", residente e domiciliado na {end}, portador do documento de identidade número ", False),
  (rg, rg_r), (", emitido por ", False), (org, org_r),
  (", inscrito no Cadastro de Pessoas Físicas do Ministério da Fazenda sob o número ", False),
  (cpf, cpf_r),
  (", na qualidade de titular dos direitos morais e patrimoniais de autor que recaem sobre minha "
   f"tese de Mestrado apresentada ao Instituto Tecnológico de Aeronáutica, intitulada “{TIT}”, com "
   "base no disposto na Lei Federal nº 9.610 de 19 de fevereiro de 1998, autorizo o Instituto "
   "Tecnológico de Aeronáutica, a partir desta data, a reproduzi-la para armazená-la permanentemente "
   "no Instituto, colocá-la ao alcance do público por meios eletrônicos, em particular mediante "
   "acesso on-line pela rede mundial de computadores, permitir a quem a ela tiver acesso que a "
   "reproduza, dela extraindo cópia, de acordo com critérios estabelecidos pelo Instituto, desde que "
   "não vise lucros, até que manifestação em sentido contrário, de minha parte, determine a cessação "
   "desta autorização.", False)]:
    r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(12)
    if isred: r.font.color.rgb = RED; r.bold = True
doc.add_paragraph()
par(f"{D['cidade']}, 31 de agosto de 2026.", 12, align=AL.RIGHT, space=36)
par("____________________________________________", 12, align=AL.RIGHT)
par("(Assinatura)", 10, align=AL.RIGHT)
doc.save(SAI / "Termo_de_Autorizacao.docx"); print("   Termo_de_Autorizacao.docx")

faltam = [k for k in ("rg_orgao","estado_civil","profissao","telefone",
                      "matricula_mes","matricula_ano") if not (D.get(k) or "").strip()]
print("\nainda faltando no JSON:", ", ".join(faltam) if faltam else "nada")
